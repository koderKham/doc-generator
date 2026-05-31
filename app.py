from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from models import db, Template, GeneratedDocument
from word_generator import WordDocumentGenerator
from docx import Document
import os
from dotenv import load_dotenv
import json
from datetime import datetime
import re

load_dotenv()

app = Flask(__name__)
CORS(app)

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOCS_FOLDER'] = 'generated_docs'
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///doc_generator.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize database
db.init_app(app)

# Create folders if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOCS_FOLDER'], exist_ok=True)

@app.before_request
def create_tables():
    """Create database tables before first request"""
    with app.app_context():
        db.create_all()

# ==================== ROUTES ====================

@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')

@app.route('/api/templates', methods=['GET'])
def get_templates():
    """Get all active templates"""
    try:
        templates = Template.query.filter_by(is_active=True).all()
        return jsonify({
            'success': True,
            'templates': [t.to_dict() for t in templates]
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/template/<int:template_id>', methods=['GET'])
def get_template(template_id):
    """Get a specific template by ID"""
    try:
        template = Template.query.get(template_id)
        if not template:
            return jsonify({'error': 'Template not found'}), 404
        
        return jsonify({
            'success': True,
            'template': template.to_dict()
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/template', methods=['POST'])
def create_template():
    """Create a new template"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data or not all(k in data for k in ['name', 'category', 'fields', 'template_content']):
            return jsonify({'error': 'Missing required fields'}), 400
        
        # Check if template already exists
        existing = Template.query.filter_by(name=data['name']).first()
        if existing:
            return jsonify({'error': 'Template with this name already exists'}), 409
        
        # Create new template
        template = Template(
            name=data['name'],
            description=data.get('description', ''),
            category=data.get('category', 'General'),
            fields=data['fields'],
            template_content=data['template_content']
        )
        
        db.session.add(template)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Template created successfully',
            'template': template.to_dict()
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/template/<int:template_id>', methods=['PUT'])
def update_template(template_id):
    """Update an existing template"""
    try:
        template = Template.query.get(template_id)
        if not template:
            return jsonify({'error': 'Template not found'}), 404
        
        data = request.get_json()
        
        # Update fields
        template.name = data.get('name', template.name)
        template.description = data.get('description', template.description)
        template.category = data.get('category', template.category)
        template.fields = data.get('fields', template.fields)
        template.template_content = data.get('template_content', template.template_content)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Template updated successfully',
            'template': template.to_dict()
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/template/<int:template_id>', methods=['DELETE'])
def delete_template(template_id):
    """Soft delete a template (mark as inactive)"""
    try:
        template = Template.query.get(template_id)
        if not template:
            return jsonify({'error': 'Template not found'}), 404
        
        template.is_active = False
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Template deleted successfully'
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload-docx-template', methods=['POST'])
def upload_docx_template():
    """Upload a Word document template and extract placeholders"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.endswith('.docx'):
            return jsonify({'error': 'Only .docx files are supported'}), 400
        
        # Save uploaded file temporarily
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(temp_path)
        
        # Extract text and find placeholders
        try:
            doc = Document(temp_path)
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Find all placeholders {{placeholder_name}}
            placeholders = re.findall(r'{{(\w+)}}', full_text)
            unique_placeholders = list(dict.fromkeys(placeholders))  # Remove duplicates, preserve order
            
            if not unique_placeholders:
                return jsonify({'error': 'No placeholders found in document. Use {{field_name}} format'}), 400
            
            # Store file for later retrieval
            template_data = {
                'filename': file.filename,
                'temp_path': temp_path,
                'placeholders': unique_placeholders,
                'full_text': full_text,
                'file_size': len(file.read())
            }
            
            return jsonify({
                'success': True,
                'message': 'Template uploaded successfully',
                'placeholders': unique_placeholders,
                'filename': file.filename,
                'preview_text': full_text[:500] + '...' if len(full_text) > 500 else full_text
            }), 200
            
        except Exception as e:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return jsonify({'error': f'Error reading document: {str(e)}'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/save-template-from-upload', methods=['POST'])
def save_template_from_upload():
    """Save uploaded template to database with field mappings"""
    try:
        data = request.get_json()
        
        if not data or not all(k in data for k in ['name', 'category', 'filename', 'fields']):
            return jsonify({'error': 'Missing required fields'}), 400
        
        # Check if template with this name already exists
        existing = Template.query.filter_by(name=data['name']).first()
        if existing:
            return jsonify({'error': 'Template with this name already exists'}), 409
        
        # Find the temporary file
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], data['filename'])
        
        if not os.path.exists(temp_path):
            return jsonify({'error': 'Uploaded file not found'}), 404
        
        # Read the document content for storage
        with open(temp_path, 'rb') as f:
            doc_content = f.read()
        
        # Store original filename for reference
        template_filename = f"template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        permanent_path = os.path.join(app.config['UPLOAD_FOLDER'], template_filename)
        
        # Move file to permanent location
        with open(permanent_path, 'wb') as f:
            f.write(doc_content)
        
        # Create template in database
        template = Template(
            name=data['name'],
            description=data.get('description', ''),
            category=data.get('category', 'General'),
            fields=data['fields'],  # Field mappings with placeholders
            template_content=permanent_path  # Store path to docx file
        )
        
        db.session.add(template)
        db.session.commit()
        
        # Clean up temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        return jsonify({
            'success': True,
            'message': 'Template saved successfully',
            'template': template.to_dict()
        }), 201
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-document', methods=['POST'])
def generate_document():
    """Generate a Word document from a template"""
    try:
        data = request.get_json()
        
        if not data or 'template_id' not in data or 'filled_data' not in data:
            return jsonify({'error': 'Missing required fields (template_id, filled_data)'}), 400
        
        template_id = data['template_id']
        filled_data = data['filled_data']
        
        # Get template
        template = Template.query.get(template_id)
        if not template:
            return jsonify({'error': 'Template not found'}), 404
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{template.name.replace(' ', '_')}_{timestamp}.docx"
        filepath = os.path.join(app.config['DOCS_FOLDER'], filename)
        
        # Check if template_content is a file path (uploaded docx) or text content
        if template.template_content.endswith('.docx') or os.path.exists(template.template_content):
            # It's an uploaded Word document
            WordDocumentGenerator.generate_from_docx(
                template.template_content,
                filled_data,
                filepath
            )
        else:
            # It's text-based template content
            WordDocumentGenerator.generate_document(
                template.template_content,
                filled_data,
                filepath
            )
        
        # Save to database
        generated_doc = GeneratedDocument(
            template_id=template.id,
            template_name=template.name,
            filename=filename,
            filled_data=filled_data,
            file_path=filepath
        )
        db.session.add(generated_doc)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Document generated successfully',
            'filename': filename,
            'download_url': f'/api/download/{generated_doc.id}'
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<int:doc_id>', methods=['GET'])
def download_document(doc_id):
    """Download a generated document"""
    try:
        doc = GeneratedDocument.query.get(doc_id)
        if not doc or not os.path.exists(doc.file_path):
            return jsonify({'error': 'Document not found'}), 404
        
        return send_file(
            doc.file_path,
            as_attachment=True,
            download_name=doc.filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get all generated documents"""
    try:
        docs = GeneratedDocument.query.order_by(GeneratedDocument.generated_at.desc()).all()
        return jsonify({
            'success': True,
            'documents': [d.to_dict() for d in docs]
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate', methods=['POST'])
def generate_docs():
    """API endpoint to generate documentation (legacy)"""
    try:
        data = request.get_json()
        
        if not data or 'content' not in data:
            return jsonify({'error': 'No content provided'}), 400
        
        content = data.get('content', '')
        doc_type = data.get('type', 'markdown')
        
        # Generate documentation based on type
        generated_doc = generate_documentation(content, doc_type)
        
        return jsonify({
            'success': True,
            'documentation': generated_doc
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def generate_documentation(content, doc_type='markdown'):
    """
    Generate documentation from input content
    
    Args:
        content (str): Input content to document
        doc_type (str): Type of documentation to generate
    
    Returns:
        str: Generated documentation
    """
    if doc_type == 'markdown':
        return generate_markdown_docs(content)
    elif doc_type == 'html':
        return generate_html_docs(content)
    else:
        return content

def generate_markdown_docs(content):
    """Generate markdown documentation"""
    lines = content.split('\n')
    docs = "# Generated Documentation\n\n"
    docs += "## Overview\n"
    docs += f"{lines[0] if lines else 'No content provided'}\n\n"
    docs += "## Details\n"
    docs += "```\n" + content + "\n```\n"
    return docs

def generate_html_docs(content):
    """Generate HTML documentation"""
    html = f"""
    <html>
    <head>
        <title>Generated Documentation</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            h1 {{ color: #333; }}
            pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; }}
        </style>
    </head>
    <body>
        <h1>Generated Documentation</h1>
        <pre>{content}</pre>
    </body>
    </html>
    """
    return html

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'healthy'}), 200

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, host='0.0.0.0', port=5000)
