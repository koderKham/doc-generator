from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class WordDocumentGenerator:
    """Generate Word documents from templates and filled data"""
    
    @staticmethod
    def generate_document(template_content, filled_data, filename):
        """
        Generate a Word document from a template and filled data
        
        Args:
            template_content (str): Template content with placeholders
            filled_data (dict): Data to fill the template
            filename (str): Output filename
        
        Returns:
            str: Path to generated document
        """
        doc = Document()
        
        # Parse template content and fill placeholders
        lines = template_content.split('\n')
        
        for line in lines:
            if line.strip():
                # Replace placeholders with actual data
                filled_line = WordDocumentGenerator._fill_placeholders(line, filled_data)
                
                # Check for formatting markers
                if filled_line.startswith('# '):
                    # Heading 1
                    heading = doc.add_heading(filled_line[2:].strip(), level=1)
                    heading.style = 'Heading 1'
                elif filled_line.startswith('## '):
                    # Heading 2
                    heading = doc.add_heading(filled_line[3:].strip(), level=2)
                    heading.style = 'Heading 2'
                elif filled_line.startswith('### '):
                    # Heading 3
                    heading = doc.add_heading(filled_line[4:].strip(), level=3)
                    heading.style = 'Heading 3'
                elif filled_line.startswith('- '):
                    # Bullet point
                    doc.add_paragraph(filled_line[2:].strip(), style='List Bullet')
                elif filled_line.startswith('* '):
                    # Numbered list
                    doc.add_paragraph(filled_line[2:].strip(), style='List Number')
                else:
                    # Regular paragraph
                    if filled_line.strip():
                        doc.add_paragraph(filled_line.strip())
        
        # Save document
        doc.save(filename)
        return filename
    
    @staticmethod
    def generate_from_docx(template_path, filled_data, output_filename):
        """
        Generate a Word document from an uploaded template (.docx)
        
        Args:
            template_path (str): Path to the template .docx file
            filled_data (dict): Data to fill the template placeholders
            output_filename (str): Output filename for generated document
        
        Returns:
            str: Path to generated document
        """
        # Load the template document
        doc = Document(template_path)
        
        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            WordDocumentGenerator._replace_text_in_paragraph(paragraph, filled_data)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        WordDocumentGenerator._replace_text_in_paragraph(paragraph, filled_data)
        
        # Save the filled document
        doc.save(output_filename)
        return output_filename
    
    @staticmethod
    def _replace_text_in_paragraph(paragraph, filled_data):
        """
        Replace placeholders in a paragraph while preserving formatting
        
        Args:
            paragraph: Paragraph object from docx
            filled_data (dict): Data to fill
        """
        # Check if paragraph contains placeholders
        full_text = paragraph.text
        placeholders = re.findall(r'{{(\w+)}}', full_text)
        
        if not placeholders:
            return
        
        # Clear existing runs and rebuild with filled data
        filled_text = WordDocumentGenerator._fill_placeholders(full_text, filled_data)
        
        # Clear paragraph
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        
        # Add filled text as new run (preserves paragraph formatting)
        if filled_text.strip():
            paragraph.add_run(filled_text)
    
    @staticmethod
    def _fill_placeholders(text, data):
        """
        Replace placeholders in text with actual data
        Placeholders format: {{field_name}}
        
        Args:
            text (str): Text with placeholders
            data (dict): Data to fill
        
        Returns:
            str: Filled text
        """
        result = text
        # Find all placeholders
        placeholders = re.findall(r'{{([^}]+)}}', text)
        
        for placeholder in placeholders:
            placeholder_key = placeholder.strip()
            # Get value from data, default to empty string
            value = data.get(placeholder_key, '')
            # Replace placeholder with value
            result = result.replace('{{' + placeholder + '}}', str(value))
        
        return result
    
    @staticmethod
    def generate_from_simple_template(template_name, filled_data, filename):
        """
        Generate a simple formatted Word document
        
        Args:
            template_name (str): Name of the template
            filled_data (dict): Data to fill
            filename (str): Output filename
        
        Returns:
            str: Path to generated document
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'{template_name}', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        from datetime import datetime
        timestamp_para = doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = timestamp_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.italic = True
        run.font.size = Pt(10)
        
        # Add content sections
        doc.add_paragraph()  # Blank line
        
        for key, value in filled_data.items():
            if value:
                # Add field name as bold
                para = doc.add_paragraph()
                run = para.add_run(f"{key.replace('_', ' ').title()}: ")
                run.bold = True
                para.add_run(str(value))
        
        # Save document
        doc.save(filename)
        return filename
